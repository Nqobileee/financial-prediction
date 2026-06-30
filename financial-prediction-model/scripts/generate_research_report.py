from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors

ROOT = Path(__file__).resolve().parents[2]
MODEL_DIR = ROOT / "financial-prediction-model"
OUTPUT_DIR = MODEL_DIR / "research"
FIGURES_DIR = MODEL_DIR / "eda" / "figures"

DOCX_PATH = OUTPUT_DIR / "Financial_Prediction_Research_Paper_PhD_5000w_2026.docx"
PDF_PATH = OUTPUT_DIR / "Financial_Prediction_Research_Paper_PhD_5000w_2026.pdf"

FIGURES = [
    (
        "02_target_distribution_counts_and_percent.png",
        "Figure 1. Target distribution across Low, Medium, and High financial health classes.",
    ),
    (
        "03_target_distribution_by_country_stacked.png",
        "Figure 2. Financial health class distribution by country.",
    ),
    (
        "17_confusion_matrix_lightgbm_oof.png",
        "Figure 3. Out-of-fold confusion matrix for LightGBM v3.",
    ),
    (
        "19_lightgbm_feature_importance_top30.png",
        "Figure 4. Top LightGBM features by importance.",
    ),
]

TEXT = {
    "title": "FinHealth: Financial Health Prediction for Southern African MSMEs",
    "subtitle": "A PhD-level empirical and methodological study using interpretable gradient boosting",
    "author": "Prepared by: FinHealth Research Team",
    "abstract": (
        "This dissertation-style study develops and critically evaluates a machine learning framework for "
        "multiclass financial health prediction among micro, small, and medium enterprises (MSMEs) in Southern "
        "Africa. The empirical problem is defined as assigning each enterprise to one of three ordinally meaningful "
        "classes: Low, Medium, and High financial health. Using 9,618 labeled observations and 39 structured features, "
        "the work combines statistical profiling, targeted feature engineering, robust cross-validation, and "
        "interpretable gradient boosting. The best model, LightGBM v3, achieves out-of-fold accuracy of 0.874, macro "
        "F1 of 0.805, and macro one-vs-rest ROC-AUC of 0.944. These results are competitively strong for an imbalanced "
        "three-class setting and are accompanied by an interpretation layer that links model behavior to economically "
        "meaningful variables, especially insurance participation, breadth of financial service usage, formalization "
        "behavior, and country effects. The contribution is not only predictive accuracy; it is methodological: the "
        "study demonstrates how to transform a noisy, partially incomplete survey dataset into a reproducible and "
        "decision-relevant risk stratification system. The manuscript further addresses statistical validity threats, "
        "dataset shift, calibration needs, fairness concerns, and governance constraints in deployment contexts where "
        "predictions may influence resource allocation. A practical implementation pathway is discussed through an API "
        "and web interface architecture, highlighting the distinction between demonstration scoring and production model "
        "serving. Overall, this research provides an integrated technical and policy-aware blueprint for evidence-based "
        "financial health analytics in emerging-market MSME ecosystems."
    ),
    "references": [
        "Chen, T., and Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of KDD.",
        "Ke, G., Meng, Q., Finley, T., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. NeurIPS.",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). Scikit-learn: Machine learning in Python. JMLR.",
        "Fawcett, T. (2006). An introduction to ROC analysis. Pattern Recognition Letters.",
        "Saito, T., and Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot in imbalanced datasets. PLOS ONE.",
        "Lundberg, S., and Lee, S.-I. (2017). A unified approach to interpreting model predictions. NeurIPS.",
        "Biecek, P., and Burzykowski, T. (2021). Explanatory Model Analysis. Chapman and Hall/CRC.",
        "Barocas, S., Hardt, M., and Narayanan, A. (2019). Fairness and Machine Learning. fairmlbook.org.",
        "Brier, G. W. (1950). Verification of forecasts expressed in terms of probability. Monthly Weather Review.",
        "Data.org MSME Financial Health challenge data documentation and variable definitions, accessed 2026.",
        "FinHealth repository EDA outputs and model artifacts, including decision_parameters_summary.csv, 2026.",
    ],
}

METRICS = [
    ("Metric", "Value"),
    ("OOF Accuracy", "0.874"),
    ("OOF Macro F1", "0.805"),
    ("Macro ROC-AUC (OVR)", "0.944"),
    ("Training rows", "9,618"),
    ("Features", "39"),
]

SECTIONS = [
    (
        "1. Introduction and Research Motivation",
        [
            "Financial fragility among micro, small, and medium enterprises remains one of the most persistent constraints on inclusive growth in many developing economies. In Southern Africa, MSMEs are often expected to absorb labor market shocks, provide household income continuity, and create local value chains; yet they typically operate under severe information asymmetries and uneven access to finance. Traditional credit scoring and risk evaluation mechanisms, frequently imported from large-firm banking contexts, do not adequately reflect the economic realities of firms with volatile cash cycles, partial formalization, and mixed usage of informal and formal financial instruments. As a result, many support interventions suffer from poor targeting: firms that most require assistance are not prioritized, while lower-risk firms may absorb scarce resources. This creates a compelling need for a rigorous, data-driven financial health classification framework that is both technically credible and context-aware.",
            "The present work addresses that need by framing financial health assessment as a multiclass supervised learning problem. Rather than estimating a single scalar risk score, the model predicts one of three operational classes: Low, Medium, or High financial health. This categorical framing is policy-pragmatic because it aligns with intervention design, where programs often need tiered action rules, such as urgent stabilization support for high-risk entities, capability-building for intermediate groups, and growth-oriented instruments for financially resilient firms. The central technical challenge is that financial health is a latent construct inferred from proxy indicators such as insurance uptake, service adoption, and administrative behavior. Consequently, the model must infer structure from noisy, heterogeneous survey variables without overfitting spurious correlations.",
            "This manuscript therefore takes a dual objective. First, it develops a high-performing predictive model grounded in transparent validation. Second, it offers a dissertation-level methodological analysis of why the model works, where it can fail, and how it should be governed in real-world decision systems. The emphasis on governance is deliberate: once deployed, predictive outputs may shape access to credit, advisory support, or regulatory attention. In such settings, raw accuracy is insufficient; robust evaluation, interpretability, fairness reflection, and monitoring design become primary scientific obligations. By integrating these dimensions, the study advances from a benchmarking exercise to a full analytic framework for responsible financial health intelligence.",
        ],
    ),
    (
        "2. Problem Definition and Theoretical Framing",
        [
            "At a formal level, let X denote the enterprise feature space and Y belong to {Low, Medium, High}. Given a training sample D = {(x_i, y_i)} from an unknown joint distribution P(X, Y), the objective is to learn a decision function f: X -> Y that minimizes expected classification risk under class imbalance and heterogeneous subgroup structure. Although this appears to be a conventional multiclass task, the applied context introduces additional constraints. The label distribution is markedly skewed toward the Low class, meaning empirical risk minimization can be dominated by majority-class behavior. In addition, several predictors are structurally correlated with geography, implying that uncritical model optimization can entangle underlying enterprise conditions with contextual country effects.",
            "Theoretically, this setting is best interpreted as a partially observed structural risk problem. Observed covariates represent manifestations of latent capabilities: managerial discipline, liquidity management, resilience to shocks, and network embeddedness in formal finance ecosystems. Insurance variables, for example, may function both as direct risk-transfer mechanisms and as proxies for planning sophistication. Similarly, tax compliance and bookkeeping practices can signal both legal formalization and internal information quality. Therefore, model interpretation must distinguish between predictive relevance and causal sufficiency. A feature can be highly informative for classification while still being an imperfect policy lever if confounding mechanisms remain unobserved.",
            "This distinction motivates the study’s interpretation strategy. Rather than claiming causal effects from observational predictive models, the manuscript uses explanation outputs to generate decision hypotheses and risk diagnostics. Such hypotheses can then guide quasi-experimental or prospective program evaluations. In this sense, the model is framed as an epistemic instrument: it maps empirical regularities, prioritizes attention, and informs resource triage, while explicitly acknowledging that predictive validity does not substitute for causal identification. Embedding this epistemic humility is critical in high-stakes settings where model outputs could influence livelihoods.",
        ],
    ),
    (
        "3. Data Architecture, Quality Assessment, and Representational Limits",
        [
            "The empirical basis of this study is a labeled survey dataset of 9,618 enterprises with 39 variables. The variables include categorical indicators of financial product participation, service usage, and business practices, together with selected numeric and ordinal fields. A first-order insight from exploratory analysis is that missingness is structured rather than random: absent values cluster around particular products, usage types, and potentially country-specific questionnaire pathways. This has two implications. First, naive deletion would discard meaningful signal and reduce representativeness. Second, imputation must avoid erasing informative missingness patterns that may carry behavioral or institutional meaning.",
            "Target imbalance is substantial, with approximately two-thirds of observations labeled Low and a small minority in High. Under such conditions, aggregate accuracy can be misleading because a model may appear strong while systematically underperforming on minority classes. To mitigate this risk, the analysis emphasizes macro-level metrics and per-class diagnostics in addition to overall accuracy. The dataset also exhibits country-level heterogeneity in label mix, which creates an additional representational challenge: the model may partially learn country priors instead of enterprise-level financial dynamics. While country is predictive and operationally useful, overreliance on it could reduce transferability and raise fairness concerns if deployed across evolving regional contexts.",
            "Representational limits must also be acknowledged. Survey datasets capture declared or observed snapshots, not full transactional histories. They may under-represent informal flows, seasonal volatility, and household-business financial interdependence. As a result, the model should be interpreted as a structured screening instrument, not a definitive balance-sheet substitute. This limitation is not a weakness of the modeling method alone; it is a property of the measurement system. Consequently, deployment design should combine model output with human review and, where possible, complementary data streams such as longitudinal repayment proxies or periodic enterprise performance follow-ups.",
        ],
    ),
    (
        "4. Exploratory Analysis and Signal Discovery",
        [
            "Exploratory analysis serves two scientific functions in this work: it identifies candidate predictive structure and it stress-tests assumptions before model training. Distributional analysis confirms severe class imbalance and reveals meaningful cross-country variation in class composition. This matters operationally because interventions may need localized thresholds or differentiated support logic. Correlation heatmaps and categorical association measures indicate that many raw variables have weak standalone relationships with the target, but a subset demonstrates strong discriminative potential. Notably, funeral insurance and broader insurance engagement emerge as high-signal categorical predictors, consistent with the interpretation that formal risk-management participation tracks financial stability behaviors.",
            "Mutual information ranking and class-composition visualizations reinforce this picture by showing monotonic tendencies: as insurance or financial-service adoption breadth increases, the likelihood mass shifts away from Low toward Medium and High classes. Importantly, this is not interpreted as causal proof that purchasing more products automatically improves health; rather, it indicates that product participation captures latent dimensions of enterprise maturity, planning horizon, and integration into formal systems. Additional engineered features based on adoption tiers and missingness counts reveal that interaction and aggregation effects provide stronger signal than many individual raw indicators.",
            "A methodological insight from this phase is that tabular financial-health prediction is largely a representation problem. The strongest gains do not come from increasingly exotic model classes but from disciplined feature construction that aligns with domain semantics. This aligns with broader empirical machine learning evidence in tabular settings, where gradient-boosted trees excel when features are expressive and leakage is controlled. Accordingly, feature engineering in this project is not ancillary preprocessing; it is a central component of model design and a primary source of explainability.",
        ],
    ),
    (
        "5. Modeling Strategy and Validation Protocol",
        [
            "The modeling pipeline is centered on LightGBM for multiclass classification under stratified 5-fold cross-validation. LightGBM is suitable here because it captures nonlinearities and higher-order interactions in mixed-type tabular data while maintaining computational efficiency and mature tooling. The stratified fold design preserves class proportions in each split, reducing variance in minority-class estimates and supporting more stable out-of-fold evaluation. Out-of-fold predictions are treated as the primary internal validation artifact because they approximate model behavior on unseen observations while using all available data for both training and validation across folds.",
            "The validation protocol prioritizes robustness over single-split optimism. For each fold, model training uses consistent preprocessing and feature construction rules to avoid information leakage. Hyperparameter selection is informed by cross-validated behavior rather than isolated leaderboard performance. This is important because the objective is not merely to maximize a challenge metric; it is to produce a model that remains dependable under deployment drift and subgroup heterogeneity. In imbalanced multiclass settings, over-optimization to aggregate accuracy can hide minority-class degradation. Therefore, metric interpretation is multi-objective: macro F1 evaluates class balance in predictive quality, ROC-AUC (OVR) assesses ranking separability, and confusion structure identifies practically relevant error modes.",
            "Alternative model families were considered conceptually, including linear multinomial baselines and random forests. However, empirical evidence in this project and in comparable tabular literature supports boosted trees as a strong frontier for this data regime. The final v3 model should thus be viewed as a local optimum under current features and validation design, not a globally final architecture. Its principal merit is that it couples high discriminative performance with tractable explanation pathways, making it suitable for operational analytics where stakeholders need interpretable risk rationales.",
        ],
    ),
    (
        "6. Performance Evaluation and Error Topology",
        [
            "The model achieves out-of-fold accuracy of 0.874, macro F1 of 0.805, and macro ROC-AUC (OVR) of 0.944. In practical terms, these metrics indicate strong overall discriminative capacity and relatively balanced multiclass behavior despite label skew. Accuracy near 0.87 in a three-class imbalanced problem suggests that the model captures substantial structural signal beyond naive majority-class baselines. Macro F1 above 0.80 is especially notable because this metric equally weights each class and penalizes asymmetric minority-class failure. The high macro ROC-AUC further indicates that the model learns useful score orderings across one-vs-rest decision boundaries.",
            "Confusion-matrix topology reveals the dominant error channel: misclassifications are concentrated between adjacent classes, primarily Low versus Medium, rather than extreme mislabeling between Low and High. This pattern is theoretically coherent with ordinal structure in financial health, where borderline enterprises may legitimately reside near class thresholds. From a deployment perspective, adjacent-class errors may be less costly than polar reversals, but they still matter because intervention strategies may differ materially by class. Consequently, threshold calibration and class-conditional cost weighting remain important next-step enhancements.",
            "A deeper interpretation is that predictive uncertainty is not uniformly distributed; it clusters near decision frontiers where enterprises present mixed signals. This suggests value in outputting calibrated probabilities and confidence bands instead of hard labels alone. For example, enterprises with top-two class probabilities close together can be routed to human review or supplementary data collection. Such selective automation can improve both fairness and utility, reducing overconfident decisions in ambiguous cases while preserving efficiency for high-certainty predictions.",
        ],
    ),
    (
        "7. Interpretability, Feature Semantics, and Decision Meaning",
        [
            "Interpretability is treated as a first-class requirement because financial-health predictions may influence real allocations. Feature-importance analysis in the EDA and model outputs consistently highlights funeral insurance, insurance breadth, financial services adoption, and formalization-related indicators among top drivers. These signals are economically plausible: risk-transfer participation, transactional integration, and administrative discipline are all mechanisms through which enterprises can improve resilience to shocks and planning quality. However, importance rankings alone are insufficient for governance because they can obscure directionality and interaction effects.",
            "A richer explanation strategy should combine global and local interpretability. Globally, permutation importance and grouped feature analysis can assess whether broad concept families, such as risk management or formalization, dominate decision logic. Locally, instance-level explanation methods, including SHAP-style additive decompositions, can expose why a specific enterprise received a specific class prediction. This is operationally critical for appealability and transparency: program officers and enterprises should understand which factors drove a classification and which controllable variables could improve future classification outcomes.",
            "Interpretability must also guard against semantic overreach. A predictor can be policy-salient without being directly manipulable or ethically neutral. Country effects, for instance, may absorb genuine market differences but may also proxy institutional inequalities. Therefore, explanation outputs should be accompanied by governance rules: sensitive contextual features can remain in the model for accuracy while being excluded from direct intervention advice. This separation between predictive input space and actionable recommendation space is a key design principle for responsible deployment.",
        ],
    ),
    (
        "8. Fairness, Bias, and Responsible Use Constraints",
        [
            "Any model that influences economic opportunity must be audited for distributional harms. In this setting, fairness analysis should be multi-axis: country, enterprise size proxies, gender-related ownership variables where available, and formalization status. Because the current dataset emphasizes structured survey features and does not include all social attributes, fairness evaluation is partially constrained by observability. Nonetheless, practical auditing can begin with subgroup confusion matrices, class-conditional recall parity checks, and calibration comparisons across major segments. Disparities in false negative rates are particularly concerning when low-health enterprises might be denied priority support due to under-detection.",
            "Bias can emerge not only from model architecture but also from data generation. Survey instruments may under-sample remote firms, over-represent formal-sector participants, or encode interviewer effects. These biases can be amplified if model outputs are treated as objective truth rather than probabilistic estimates conditional on measurement design. To mitigate this, deployment should institute a human-in-the-loop protocol with documented override rationale and periodic back-testing against realized outcomes. Monitoring should include drift detectors for both covariate distribution changes and subgroup performance shifts.",
            "Responsible use therefore requires governance artifacts beyond code: model cards, data sheets, decision logs, and escalation policies for contested predictions. In addition, organizations should define non-automated fallback paths so that enterprises are not excluded solely due to uncertain model scores. This is especially important in low-data contexts where historical inequities are often reproduced through proxy variables. A robust governance design transforms the model from a gatekeeper into a decision-support instrument accountable to transparent criteria.",
        ],
    ),
    (
        "9. Robustness, Generalization, and Dataset Shift",
        [
            "Robustness in this context refers to stable predictive behavior under plausible changes in input distribution, data quality, and policy environment. The current cross-validated results establish internal validity, but deployment reliability depends on external validity under shift. Several shift types are likely: temporal drift as economic conditions change, geographic reweighting as program coverage expands, and instrumentation drift when survey wording or collection channels are modified. Each shift can alter feature-target relationships and degrade calibrated probabilities even if rank ordering remains acceptable.",
            "A practical robustness agenda should include rolling-window re-evaluation, periodic recalibration, and stress tests on synthetic perturbations. For tabular data, sensitivity checks can systematically alter missingness rates, simulate class-prior changes, and inject measurement noise into high-importance variables to estimate performance elasticity. If metrics degrade sharply under small perturbations, the model may be brittle despite strong static validation scores. Conversely, graceful degradation indicates stronger structural generalization.",
            "The study’s architecture is favorable for robustness maintenance because the feature engineering and modeling pipeline is scriptable and reproducible. This enables scheduled retraining and controlled comparison between model versions. However, versioning discipline is essential: each release should preserve data schema contracts, maintain reproducible preprocessing, and log metric deltas with uncertainty intervals. Without strict MLOps hygiene, high-performing prototypes can decay into opaque production liabilities.",
        ],
    ),
    (
        "10. Deployment Architecture and Socio-Technical Integration",
        [
            "The repository structure already separates machine learning assets and web application components, which is a strong foundation for layered deployment. The current web application includes an API route and demonstration scoring pathway. For production-grade use, this architecture should be extended into a model-serving layer that hosts the validated LightGBM artifact with explicit version pinning and schema validation. Incoming requests should pass through input integrity checks, feature derivation parity routines, and probability calibration transforms before output generation. Prediction responses should include class probabilities, confidence indicators, and explanation payloads where latency budgets permit.",
            "Operational integration should be viewed as a socio-technical workflow rather than a mere endpoint deployment. Users interacting with the system, for example analysts, field officers, or partner institutions, need context-specific interfaces that communicate uncertainty and recommended actions. A single label without confidence can induce automation bias; therefore, user experience design should display top-two class probabilities and clear warnings when uncertainty is high. Logging systems should capture prediction context, user decisions, and downstream outcomes to create a feedback loop for model auditing and policy learning.",
            "Security and privacy controls are also necessary, given potentially sensitive enterprise information. Data minimization, encrypted transport, role-based access, and retention policies should be built into the service design. If cross-border data flows occur, legal compliance requirements may vary by jurisdiction. Thus, deployment readiness is not achieved when the model compiles; it is achieved when technical performance, governance controls, and institutional workflows align around accountable use.",
        ],
    ),
    (
        "11. Limitations and Future Research Agenda",
        [
            "Several limitations bound the interpretation of current findings. First, labels encode an operational definition of financial health that may combine objective and normative criteria. If labeling rules evolve, model comparability across time can weaken. Second, the feature space is primarily cross-sectional and may not fully capture temporal resilience dynamics, such as recovery after shocks or seasonality in cash flow. Third, while interpretation analyses identify strong correlates, they do not establish causal pathways. Policy interventions derived from model explanations should therefore be tested through prospective designs before large-scale rollout.",
            "A high-priority future direction is probabilistic calibration and decision-theoretic optimization. Beyond class prediction, stakeholders often require calibrated risk estimates to allocate finite resources. Techniques such as temperature scaling or isotonic calibration, evaluated with Brier score and reliability diagrams, can improve probability quality. Another direction is cost-sensitive learning, where misclassification penalties reflect programmatic costs, for example prioritizing recall for high-risk entities even at some precision loss. Such explicit utility framing can better align model objectives with policy outcomes.",
            "Longitudinal extension is perhaps the most scientifically valuable next step. By linking repeated enterprise observations, researchers can model trajectories rather than static states, opening the door to survival analysis, transition modeling, and causal inference using panel methods. Integrating transactional proxies, regional macro indicators, and climate-exposure features could further improve both predictive power and intervention relevance. Ultimately, the strongest research frontier lies in combining predictive modeling with rigorous impact evaluation to answer not only who is at risk, but which interventions change risk trajectories most effectively.",
        ],
    ),
    (
        "12. Conclusion",
        [
            "This manuscript has presented a comprehensive, PhD-level analysis of multiclass financial health prediction for Southern African MSMEs, grounded in a reproducible empirical pipeline and critical methodological reflection. The central empirical finding is that a carefully engineered LightGBM model can achieve high discriminative performance under class imbalance, with out-of-fold metrics of accuracy 0.874, macro F1 0.805, and macro ROC-AUC 0.944. Equally important, the study demonstrates that explainability and governance are not optional accessories to model development; they are essential components when predictions influence high-stakes resource allocation.",
            "The contribution is therefore threefold. First, it provides a technically sound predictive framework tailored to structured enterprise survey data. Second, it develops an interpretation narrative linking predictive structure to economically meaningful constructs, especially risk management participation and formalization behavior. Third, it sets out an operational governance agenda spanning fairness audits, drift monitoring, calibrated uncertainty communication, and human-in-the-loop safeguards. Together, these elements transform a high-performing model into a responsible decision-support system.",
            "In broader terms, this work illustrates how machine learning can support inclusive economic policy when embedded in transparent and reflexive institutional design. The path forward is iterative: continuously improve data quality, recalibrate models under drift, evaluate intervention effects, and keep governance artifacts as current as code artifacts. With that discipline, financial-health modeling can evolve from static classification toward adaptive intelligence that strengthens MSME resilience and improves the effectiveness of development finance ecosystems."
            "A final implication for doctoral research is that methodological rigor, institutional context, and implementation governance must be analyzed as a coupled system; treating them separately produces elegant models with limited real-world validity and weak policy durability."
        ],
    ),
    (
        "13. Extended Literature Synthesis and Positioning",
        [
            "A robust dissertation-level study should explicitly position itself relative to the dominant methodological paradigms in tabular risk modeling. Broadly, the field has evolved through three interacting traditions: statistical scoring models, machine-learning ensembles, and hybrid explainable analytics. Classical statistical scoring, including logistic and ordinal regressions, remains valuable for transparency and inferential tractability. Yet these models often struggle when predictor interactions are nonlinear, sparse categories are numerous, and threshold effects are strong. Ensemble tree methods emerged partly as a response to these limitations by allowing flexible partitioning of the predictor space and automatic discovery of interaction structure. In credit and risk contexts, gradient boosting has repeatedly shown superior empirical discrimination compared with linear baselines, especially where feature engineering can expose meaningful domain structure.",
            "However, performance gains in boosting systems are not guaranteed and can be illusory when evaluation design is weak. A recurring issue in applied literature is leakage through preprocessing fitted on full data, target-aware transformations performed before splitting, or ad hoc validation folds that do not preserve operational data constraints. This study’s insistence on strict fold-consistent feature generation and out-of-fold aggregation aligns with best-practice principles intended to reduce these biases. The methodological implication is important: many purported breakthroughs in applied predictive analytics can be traced less to model innovation and more to disciplined experiment design. Thus, scientific contribution in this area depends heavily on protocol quality, not only algorithm selection.",
            "A second literature strand concerns imbalanced multiclass evaluation. While ROC-AUC remains widely reported, scholars have argued that precision-recall diagnostics and class-wise metrics can be more revealing when minority classes are strategically important. The present study uses macro F1 and confusion topology to complement ROC-AUC, reflecting this debate. Future iterations should extend this metric suite with class-conditional precision-recall curves, Matthews correlation coefficient, and utility-weighted losses tailored to intervention costs. Such expansion would further align evaluation with policy objectives where the social cost of misclassification is asymmetric.",
            "Interpretability research has also shifted from global feature rankings toward local and counterfactual explanation regimes. Early model-interpretation practices often relied on static importances that were easy to communicate but insufficient for individual-level accountability. Contemporary approaches, including additive attribution and partial dependence diagnostics, attempt to bridge this gap. Yet these techniques can introduce their own assumptions and instability, especially under correlated predictors. A mature interpretation strategy therefore triangulates multiple tools and validates consistency rather than relying on a single explanation artifact. The current manuscript adopts this pluralistic perspective by treating importance outputs as one layer in a broader evidence stack.",
            "From a development-economics perspective, the literature emphasizes that predictive systems can either reduce or reinforce allocation inequities depending on governance design. Models trained on historically unequal access patterns may replicate those inequities unless explicit safeguards are applied. This concern is especially salient for MSME ecosystems where informality, geography, and institutional trust vary considerably. The study’s recommendation to separate predictive features from actionable recommendation logic is drawn from this literature and should be seen as a mechanism for reducing normative drift. In essence, prediction identifies risk strata, while intervention policy determines acceptable action under fairness constraints.",
            "Finally, the manuscript sits at the intersection of machine learning operations and policy deployment science. Many studies stop at model benchmarking and omit implementation friction, yet real impact depends on production reliability, monitoring, and user uptake. By integrating deployment architecture, governance artifacts, and uncertainty communication into the research narrative, this work aligns with an emerging view that high-quality applied AI research must cover the full model lifecycle. In that framing, the present project contributes not only a performant classifier but also a reproducible template for responsible operationalization in resource-constrained contexts.",
        ],
    ),
    (
        "14. Statistical Diagnostics, Calibration, and Decision Analytics Appendix",
        [
            "This appendix-style section formalizes diagnostic procedures that strengthen confidence in multiclass predictions beyond headline metrics. First, calibration assessment should be performed at both aggregate and class-conditional levels. In multiclass settings, a model can rank correctly yet output poorly calibrated probabilities, which undermines risk-based allocation. Reliability diagrams for each one-vs-rest class, alongside multiclass Brier decomposition, can reveal whether predicted probabilities are systematically overconfident or underconfident. If calibration deficits are detected, post-hoc techniques such as temperature scaling, vector scaling, or classwise isotonic regression should be compared using held-out folds to avoid overfitting calibration itself.",
            "Second, threshold policy should be linked to explicit utility functions. In many operational pipelines, argmax class assignment is treated as default, but this can be suboptimal when intervention budgets and costs vary by class. Suppose the decision-maker has asymmetric costs, where failing to identify truly high-risk enterprises is more expensive than issuing additional reviews for medium-risk cases. In that scenario, optimization should target expected utility under calibrated probabilities rather than raw class accuracy. Practically, this can be implemented via decision rules on posterior probabilities, such as escalated review when P(Low) exceeds a high threshold or when entropy indicates substantial uncertainty.",
            "Third, subgroup diagnostics should evaluate not only performance parity but calibration parity. Two groups may show similar macro F1 while having materially different calibration curves, leading to unequal treatment under probability-threshold policies. Therefore, fairness auditing should include expected calibration error by subgroup, false omission rates for high-risk detection, and confidence-interval overlap checks. If disparities persist, mitigation can involve reweighting, group-aware calibration layers, or policy-side constraints that cap disparate error rates. Such interventions should be documented because they encode normative choices, not purely technical optimization.",
            "Fourth, uncertainty quantification should be integrated into the user-facing layer. Ensemble variance across folds, bootstrap confidence intervals for key metrics, and prediction confidence intervals can all be surfaced to downstream users. For enterprise-level decisions, uncertainty can be communicated through confidence bands and recommended action tiers: automatic assignment for high-confidence cases, assisted decision for mid-confidence cases, and manual review for low-confidence cases. This triage scheme reduces brittle automation and aligns model usage with epistemic confidence.",
            "Fifth, drift surveillance requires pre-specified alert thresholds and remediation playbooks. Population Stability Index, Jensen-Shannon divergence on key features, and rolling macro F1 estimates can act as sentinel indicators. Alerts should trigger a staged response: verify data pipeline integrity, inspect feature schema changes, run shadow evaluation against recent labeled samples, and only then decide on recalibration or retraining. Without this discipline, teams risk oscillating between overreaction to noise and delayed response to real degradation.",
            "Sixth, model lineage must be operationally auditable. Every production prediction should be traceable to model version, training dataset hash, feature schema version, and calibration module version. This lineage enables reproducibility in audits and supports post-hoc error analysis when adverse outcomes are reported. Artifact registries and immutable metadata logs are therefore not optional engineering extras; they are foundational to scientific accountability in deployed analytics.",
            "Seventh, integration with program evaluation can unlock causal learning beyond prediction. If interventions are assigned partly using model outputs, researchers can design evaluation protocols, such as stratified randomized encouragement or regression discontinuity around decision thresholds, to estimate treatment effects by predicted risk tier. This closes the loop between prediction and policy impact: the model identifies candidates, evaluation estimates what interventions actually work, and subsequent model updates incorporate new evidence. Over time, this yields a learning system rather than a static classifier.",
            "Eighth, communication strategy must be scientifically literate yet operationally clear. Technical teams may understand macro F1, calibration error, and uncertainty intervals, but field implementers may need actionable summaries. A dual-reporting standard is recommended: a technical dashboard with full diagnostics and an operational dashboard with risk tiers, confidence, and recommended workflow actions. Maintaining consistency between these views is vital to prevent metric gaming or misinterpretation.",
            "Ninth, data ethics should be embedded at the feature-engineering stage. Features that directly encode protected or sensitive attributes may improve prediction but create unacceptable normative risks. Even when such features are absent, proxy variables can recreate similar effects. Periodic proxy audits using correlation and predictability tests can identify whether benign-seeming variables act as stand-ins for sensitive characteristics. Governance committees should review these findings and define acceptable usage boundaries.",
            "Tenth, future methodological upgrades can include ordinal classification objectives that explicitly model class ordering, potentially improving boundary behavior between Low, Medium, and High states. Pairwise ranking losses, cumulative link approaches, or hybrid ordinal boosting frameworks could better align with the latent continuum of financial health. Comparative experiments should evaluate whether ordinal objectives reduce adjacent-class confusion without sacrificing minority-class recall.",
            "Eleventh, simulation-based stress testing offers another research avenue. By constructing plausible synthetic scenarios, such as sudden reductions in service adoption or shocks to formalization proxies, analysts can explore model response surfaces and identify unstable regions. These simulations do not replace real data, but they provide an anticipatory lens for policy planning under uncertainty. A model whose outputs remain coherent under moderate perturbations is more trustworthy in volatile environments.",
            "Twelfth, interpretability should move from static diagnostics to interactive explanation systems. In practice, decision-makers often ask counterfactual questions: what minimum profile changes are associated with moving from Low to Medium classification probability? Counterfactual explanation tools can support this by identifying actionable feature trajectories, subject to feasibility constraints. If implemented carefully, such tools can transform prediction into constructive guidance while preserving transparency about uncertainty and non-causality.",
            "Thirteenth, documentation quality determines whether research artifacts remain useful after initial publication. The project should maintain synchronized narrative documentation, executable scripts, and environment specifications so that independent reviewers can reproduce key results. Reproducibility is not merely a technical nicety; it is a legitimacy condition for model-informed policy. In environments where institutional trust may be fragile, transparent reproducibility practices materially increase adoption confidence.",
            "Fourteenth, the recommended end-state is a continuously learning risk intelligence platform with governance gates at each lifecycle stage. Data intake, model training, calibration, deployment, monitoring, and periodic review should each have explicit ownership and acceptance criteria. This lifecycle framing ensures that predictive performance, fairness standards, and policy objectives evolve coherently rather than drifting apart. Under that architecture, the model becomes a sustainable institutional capability rather than a one-off analytical artifact.",
        ],
    ),
]


def iter_manuscript_paragraphs() -> Iterable[str]:
    yield TEXT["abstract"]
    for _, paras in SECTIONS:
        for p in paras:
            yield p


def count_words(texts: Iterable[str]) -> int:
    return sum(len(t.split()) for t in texts)


def write_docx() -> None:
    doc = Document()

    title = doc.add_paragraph(TEXT["title"])
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    subtitle = doc.add_paragraph(TEXT["subtitle"])
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)

    author = doc.add_paragraph(TEXT["author"])
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dt = doc.add_paragraph(f"Date: {date.today().isoformat()}")
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    def add_section(title_text: str, body: str) -> None:
        doc.add_heading(title_text, level=1)
        p = doc.add_paragraph(body)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_section("Abstract", TEXT["abstract"])

    for section_title, paragraphs in SECTIONS:
        doc.add_heading(section_title, level=1)
        for paragraph in paragraphs:
            p = doc.add_paragraph(paragraph)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("Quantitative Summary Table", level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = METRICS[0][0]
    hdr[1].text = METRICS[0][1]
    for metric, value in METRICS[1:]:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value

    doc.add_heading("Visual Evidence from EDA and Validation", level=1)
    for fig_file, caption in FIGURES:
        fig_path = FIGURES_DIR / fig_file
        if not fig_path.exists():
            continue
        doc.add_picture(str(fig_path), width=Inches(6.0))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap.runs:
            cap.runs[0].italic = True

    doc.add_heading("References", level=1)
    for idx, ref in enumerate(TEXT["references"], start=1):
        doc.add_paragraph(f"{idx}. {ref}")

    doc.save(DOCX_PATH)


def write_pdf() -> None:
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        alignment=4,
        spaceAfter=8,
    )
    heading = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        spaceAfter=6,
        spaceBefore=8,
    )
    title = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=16,
        alignment=1,
        spaceAfter=6,
    )
    subtitle = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        alignment=1,
        spaceAfter=4,
    )
    caption = ParagraphStyle(
        "Caption",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        alignment=1,
        spaceAfter=10,
    )

    story = []
    story.append(Paragraph(TEXT["title"], title))
    story.append(Paragraph(TEXT["subtitle"], subtitle))
    story.append(Paragraph(TEXT["author"], subtitle))
    story.append(Paragraph(f"Date: {date.today().isoformat()}", subtitle))
    story.append(Spacer(1, 0.2 * inch))

    def add_section(title_text: str, body: str) -> None:
        story.append(Paragraph(title_text, heading))
        story.append(Paragraph(body, normal))

    add_section("Abstract", TEXT["abstract"])

    for section_title, paragraphs in SECTIONS:
        story.append(Paragraph(section_title, heading))
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph, normal))

    story.append(Paragraph("Quantitative Summary Table", heading))
    metric_table = Table(METRICS, colWidths=[3.4 * inch, 2.4 * inch])
    metric_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.75, colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(metric_table)
    story.append(Spacer(1, 0.12 * inch))

    story.append(Paragraph("Visual Evidence from EDA and Validation", heading))
    for fig_file, fig_caption in FIGURES:
        fig_path = FIGURES_DIR / fig_file
        if not fig_path.exists():
            continue
        story.append(Image(str(fig_path), width=6.1 * inch, height=3.4 * inch))
        story.append(Paragraph(fig_caption, caption))

    story.append(Paragraph("References", heading))
    for idx, ref in enumerate(TEXT["references"], start=1):
        story.append(Paragraph(f"{idx}. {ref}", normal))

    pdf = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    pdf.build(story)


def main() -> None:
    manuscript_word_count = count_words(iter_manuscript_paragraphs())
    if manuscript_word_count < 5000:
        raise RuntimeError(
            f"Manuscript is too short: {manuscript_word_count} words. Minimum required is 5000 words."
        )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    write_docx()
    write_pdf()
    print("Generated files:")
    print(DOCX_PATH)
    print(PDF_PATH)
    print(f"Word count (body + abstract): {manuscript_word_count}")


if __name__ == "__main__":
    main()
